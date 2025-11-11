# made by nickson
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import os
import re
import datetime

INT_TO_MONTHS = {
    1 : "Januari",
    2 : "Februari",
    3 : "Maret",
    4 : "April",
    5 : "Mei",
    6 : "Juni",
    7 : "Juli",
    8 : "Agustus",
    9 : "September",
    10 : "Oktober",
    11 : "November",
    12 : "Desember" 
}

class Surat():

    def __init__(self, path="", directory_name="", NoSur='', due='', theme='', alkitab='', Fokus=[], NamPF = "", temaItalic = False, SuratDalam = False, mon="", years=""):
        if path == '' or not os.path.exists(path) or not os.path.isfile(path): raise ValueError("Path must not be blank and exist as a file")
        if directory_name == '' or not os.path.exists(directory_name) or not os.path.isdir(directory_name): raise ValueError("Directory name must not be blank and exist as a dir")
        if not SuratDalam:
            if not (NoSur and NamPF and due and theme and alkitab and Fokus) : raise ValueError("Parms can't be blank")
        elif SuratDalam:
            if not (NoSur and due and theme and alkitab and Fokus) : raise ValueError("Parms can't be blank")

        self._path = path
        self._directory_name = directory_name
        self._NoSur = NoSur
        self._NamPF = NamPF
        self._due = due
        self._theme = theme
        self._alkitab = alkitab
        self._Fokus = Fokus
        self._temaItalic = temaItalic
        self._SuratDalam = SuratDalam

        dir, input_name = os.path.split(path)
        dir, curr_fol = os.path.split(dir)
        input_name, exe = os.path.splitext(input_name)
        input_name = " ".join(input_name.split(" ")[:-1])
        output_name = input_name+ ' ' + due + exe

        self._dir = dir
        self._curr_fol = curr_fol
        self._output_name = output_name

        self._finding = {
            "Sysdate": self.curr_time(), 
            "NoSurat": str(NoSur), 
            "SysMon": self.Romans(mon), 
            "SysYear": str(years), 
            "NamaPF": NamPF, 
            "DueDate": due, 
            "Theme": theme, 
            "BacaanAlkitab": alkitab}

    
    def make_document(self):
        self._document_read = docx.Document(self._path)

    def add_text(self):
        for param in self._document_read.paragraphs:
            for run in param.runs:
                if run.text == '[' or run.text == ']': run.text = ''
                
                mathes = re.findall(r'\[(.*?)\]', run.text)
                if mathes:
                    for k in mathes:
                        if k == "FokusTema":
                            run.text = self._Fokus[0]

                            for foc in self._Fokus[1:]:
                                new_para = param.insert_paragraph_before('')
                                new_para.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE
                                self.set_bullet(new_para)

                                new_run = new_para.add_run(foc)
                                new_run.font.italic = True
                                new_run.font.name = 'Times New Roman'
                                new_run.font.size = docx.shared.Pt(12)
                                
                            break
                        if k == "Theme":
                            run.font.italic = self._temaItalic
                        run.text = run.text.replace(f"[{k}]", self._finding[k])
                elif run.text in self._finding:
                    run.text = run.text.replace(f"{run.text}", self._finding[run.text])
                elif run.text == "FokusTema":
                    run.text = self._Fokus[0]
                    
                    param.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE
                    for foc in self._Fokus[1:]:
                        new_para = param.insert_paragraph_before('')
                        new_para.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.ONE_POINT_FIVE
                        self.set_bullet(new_para)

                        new_run = new_para.add_run(foc)
                        new_run.font.italic = True
                        new_run.font.name = 'Times New Roman'
                        new_run.font.size = docx.shared.Pt(12)


    def save_document(self):
        try:
            self._document_read.save(os.path.join(self._directory_name, self._output_name) )
        except FileExistsError:
            raise FileExistsError(f"The name {os.path.join(self._directory_name, self._output_name)} has a file")

    def Romans(self, mon):
        roman = {
                1 : "I",
                4 : "IV",
                5 : "V",
                9 : "IX",
                10 : "X",
                40 : "XL",
                50 : "L",
                90 : "XC",
                100 : "C",
                400 : "CD",
                500 : "D",
                900 : "CM",
                1000 : "M"
            }
        num = mon
        num = str(num)
        power = len(num) - 1
        re = ""
        for i in range(len(num)):
            cur = 10**power
            rom = int(num[i]) 
            power -= 1
            # print(cur)
            if num[i] == "9" or num[i] == "4":
                re += roman[rom*cur]
            elif rom >= 5:
                re += roman[cur*5]
                re += roman[cur]*(rom-5)
            else:
                re += roman[cur]*rom
        return re

    def curr_time(self) -> str:
        t = datetime.date.today()
        return f'{t.day} {INT_TO_MONTHS[t.month]} {t.year}'
    
    # Making the bullet list
    def set_bullet(self, paragraph, level=0):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')  # assumes default bullet numbering
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

# Deleting the dir, iff there is no dir
def delete_dir(dir_path):
    
    for filename in os.listdir(dir_path):
        file_path = os.path.join(dir_path, filename)
        if os.path.isdir(file_path):
            print(f"Failed to delete Directory {os.path.split(dir_path)[1]}, because there is a directory inside ")
            raise ValueError("There is a directory inside")

    for filename in os.listdir(dir_path):
        file_path = os.path.join(dir_path, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)
            print(f"Deleted file: {filename}")

    try:
        os.rmdir(dir_path)
    except OSError as Error:
        print(Error)
        print(f"Directory {os.path.split(dir_path)[1]} can not be removed")

# make a dir iff there is no file or dir that has the same name on the file 
def make_dir(path, s_mon,years):
    directory_name = s_mon + ' ' + years
    paths = os.path.join(path,directory_name)
    try:
        os.mkdir(paths)
        print(f"Directory '{directory_name}' created.")
    except FileExistsError:
        print(f"Directory '{directory_name}' already exists.")
        return 1
    except OSError as e:
        print(f"Error creating directory: {e}")
        return -1
    
    return paths

# Zipping the content of the dir
def zipping(paths, Send_to):
    def zipdir(path, ziph):
        # ziph is zipfile handle
        for root, dirs, files in os.walk(path):
            for file in files:
                ziph.write(os.path.join(root, file), 
                        os.path.relpath(os.path.join(root, file), 
                                        os.path.join(path, '..')))
        


    with zipfile.ZipFile(f'{Send_to}.zip', 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipdir(paths, zipf)
        print("File has Been Zipped!")


if __name__ == "__main__":
    # Usage
    # Making the dir
    name_dir = make_dir("/mnt/c/Coding Folder/Python/Surat_Pdt/", "Desember","2025")
    # Making the Surat
    Surat_Luar = Surat(path = "static/Surat Pendeta Luar.docx",
                       directory_name=name_dir, 
                       NoSur='47', 
                       due="30 November 2025", 
                       theme="Together for Change", 
                       temaItalic=True, 
                       alkitab="Matius 24:36-44", 
                       Fokus=['Remaja mengerti bahwa berjaga-jaga untuk kedatangan Yesus kedua kali juga berarti membawa perubahan yang berdampak baik bagi dunia.','Remaja diajak untuk bersama-sama menciptakan perubahan baik di lingkungannya, mulai dari hal-hal kecil.'], 
                       NamPF="Pdt. Em. Samuel Santoso (GKI Kedoya)", 
                       years="2025", 
                       mon="12")
    # Making the doc
    Surat_Luar.make_document()
    # removing placeholder and adding text
    Surat_Luar.add_text()
    # saving in the dir, with the format name Surat Pendeta [due].docx
    Surat_Luar.save_document()


    Surat_Dalem = Surat(path="static/Surat Pendeta Pak_Naya.docx", directory_name=name_dir, NoSur='44', due='9 November 2025', theme='Heroes of Faith', temaItalic=True, alkitab='1 Timotius 4:12-16', SuratDalam=True, Fokus=["Remaja diajak untuk belajar melihat dan melatih dirinya untuk menjawab panggilan Tuhan sebagai pahlawan iman di tengah tantangan dunia.","Remaja berkomitmen untuk menjadi “pahlawan” dengan memberikan teladan yang baik."], years="2025", mon="12")
    Surat_Dalem.make_document()
    Surat_Dalem.add_text()
    Surat_Dalem.save_document()
    
    Surat_Dalem = Surat(path="static/Surat Pendeta Kak_Angela.docx", directory_name=name_dir, NoSur='45', due='16 November 2025', theme='Heroes of Faith', temaItalic=True, alkitab='1 Timotius 4:12-16', SuratDalam=True, Fokus=["Remaja diajak untuk belajar melihat dan melatih dirinya untuk menjawab panggilan Tuhan sebagai pahlawan iman di tengah tantangan dunia.","Remaja berkomitmen untuk menjadi “pahlawan” dengan memberikan teladan yang baik."],years="2025", mon="12")
    Surat_Dalem.make_document()
    Surat_Dalem.add_text()
    Surat_Dalem.save_document()
    
    Surat_Dalem = Surat(path="static/Surat Pendeta Kak_Gloria.docx", directory_name=name_dir, NoSur='46', due='23 November 2025', theme='Heroes of Faith', temaItalic=True, alkitab='1 Timotius 4:12-16', SuratDalam=True, Fokus=["Remaja diajak untuk belajar melihat dan melatih dirinya untuk menjawab panggilan Tuhan sebagai pahlawan iman di tengah tantangan dunia.","Remaja berkomitmen untuk menjadi “pahlawan” dengan memberikan teladan yang baik."], years="2025", mon="12")
    Surat_Dalem.make_document()
    Surat_Dalem.add_text()
    Surat_Dalem.save_document()

    # Zipping the docx content in dir
    zipping(name_dir, name_dir)
    # deleting dir
    delete_dir(name_dir)